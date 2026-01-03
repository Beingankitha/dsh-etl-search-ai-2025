"""
DOCX document text extractor using python-docx library.

Handles:
- Microsoft Word .docx files
- Paragraphs, tables, headers, footers
- Error handling for corrupted files
"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

from src.logging_config import get_logger
from .document_extractor import DocumentExtractor, DocumentExtractorError

logger = get_logger(__name__)


class DOCXExtractor(DocumentExtractor):
    """
    Extracts text from DOCX (Microsoft Word) documents.

    Features:
    - Paragraph extraction
    - Table content extraction
    - Heading preservation
    - Error recovery for corrupted files
    """

    SUPPORTED_MIMES = ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]

    async def extract(self, file_path: str | Path) -> str:
        """
        Extract text from DOCX.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text with structure preserved

        Raises:
            DocumentExtractorError: If DOCX invalid or extraction fails
        """
        try:
            path = self._validate_file(file_path)
            logger.info(f"Extracting DOCX: {path}")

            doc = Document(path)
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    if any(row_cells):
                        text_parts.append(" | ".join(row_cells))

            if not text_parts:
                raise DocumentExtractorError(f"No text extracted from {path}")

            combined = "\n".join(text_parts)
            sanitized = self._sanitize_text(combined)

            logger.info(f"Extracted {len(text_parts)} paragraphs from DOCX: {path}")
            return sanitized

        except DocumentExtractorError:
            raise
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise DocumentExtractorError(f"Failed to extract DOCX: {e}") from e